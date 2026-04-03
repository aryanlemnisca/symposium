"""Document text extraction — PDF, DOCX, XLSX, CSV, TXT, MD."""

import io
import os


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from a document. Raises ValueError on unsupported/failed extraction."""
    ext = os.path.splitext(filename)[1].lower()

    if ext in (".txt", ".md"):
        return file_bytes.decode("utf-8", errors="replace")

    if ext == ".pdf":
        return _extract_pdf(file_bytes, filename)

    if ext == ".docx":
        return _extract_docx(file_bytes)

    if ext == ".xlsx":
        return _extract_xlsx(file_bytes)

    if ext == ".csv":
        return _extract_csv(file_bytes)

    raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_bytes: bytes, filename: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    has_images = False
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
        if hasattr(page, "images") and len(page.images) > 0:
            has_images = True

    result = "\n\n".join(pages).strip()
    if not result:
        raise ValueError(f"Could not extract text from {filename}. The PDF may be image-only.")
    if has_images:
        result += (
            "\n\n[Note: This document contains images/charts that could not be "
            "extracted as text. Key visual content should be described in the problem statement.]"
        )
    return result


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, header_sep)
            parts.append("\n".join(rows))

    return "\n\n".join(parts).strip()


def _extract_xlsx(file_bytes: bytes) -> str:
    import pandas as pd

    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    parts = []
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        parts.append(f"## Sheet: {sheet_name}\n\n{df.to_markdown(index=False)}")
    return "\n\n".join(parts).strip()


def _extract_csv(file_bytes: bytes) -> str:
    import pandas as pd

    df = pd.read_csv(io.BytesIO(file_bytes))
    return df.to_markdown(index=False)
